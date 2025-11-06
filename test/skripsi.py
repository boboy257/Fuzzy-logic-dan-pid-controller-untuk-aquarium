from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFont
import os

# === FUNGSI: BUAT PLACEHOLDER GAMBAR ===
def buat_placeholder(teks, path):
    img = Image.new('RGB', (600, 300), color=(240, 240, 240))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.load_default()
    except:
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), teks, font=font)
    x = (600 - (bbox[2] - bbox[0])) // 2
    y = (300 - (bbox[3] - bbox[1])) // 2
    draw.text((x, y), teks, fill=(0, 0, 0), font=font)
    img.save(path)

# Buat placeholder jika belum ada
gambar_list = [
    ("gambar1_iot.png", "Gambar 1. Arsitektur Sistem IoT pada Monitoring Kualitas Air Akuarium"),
    ("gambar2_pid.png", "Gambar 2. Kurva Respons PID Controller terhadap Gangguan Suhu"),
    ("gambar3_fuzzy.png", "Gambar 3. Blok Diagram Fuzzy Logic Controller"),
    ("gambar4_eval.png", "Gambar 4. Parameter Evaluasi Sistem Kendali")
]
for path, teks in gambar_list:
    if not os.path.exists(path):
        buat_placeholder(teks, path)

# === BUAT DOKUMEN ===
doc = Document()

# Gaya dasar
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def h1(teks): doc.add_heading(teks, level=1)
def h2(teks): doc.add_heading(teks, level=2)
def p(teks): doc.add_paragraph(teks)
def img(path): 
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(5.0))
def caption(teks):
    para = doc.add_paragraph()
    para.add_run(teks).italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ========================
# HALAMAN JUDUL
# ========================
p("ANALISIS PERBANDINGAN PID CONTROLLER DAN FUZZY LOGIC PADA MONITORING KUALITAS AIR AKUARIUM BERBASIS")
p("INTERNET OF THINGS")
p("")
p("HAFIZ MOHAMMAD ISKANDAR")
p("D121 19 10 12")
p("")
p("PROGRAM STUDI SARJANA TEKNIK INFORMATIKA")
p("FAKULTAS TEKNIK")
p("UNIVERSITAS HASANUDDIN")
p("GOWA")
p("2025")
doc.add_page_break()

# ========================
# ABSTRAK
# ========================
h1("ABSTRAK")
p("Hafiz Mohammad Iskandar. Analisis Perbandingan PID Controller dan Fuzzy Logic pada Monitoring Kualitas Air Akuarium Berbasis Internet of Things. (Dibimbing oleh Muhammad Alief Fahdal Imran Oemar)")
p("")
p("Pemeliharaan akuarium membutuhkan stabilitas parameter kualitas air seperti suhu dan kekeruhan (turbidity) agar ikan tetap sehat. Metode manual yang umum digunakan seringkali tidak akurat dan rentan terhadap kesalahan manusia. Penelitian ini bertujuan untuk merancang sistem monitoring kualitas air akuarium berbasis Internet of Things (IoT) yang menerapkan dua metode kendali: PID Controller dan Fuzzy Logic, serta membandingkan performa keduanya berdasarkan parameter overshoot, steady-state error, dan settling time.")
p("")
p("Sistem dirancang menggunakan sensor DS18B20 untuk suhu dan sensor turbidity AB147 untuk mengukur kekeruhan air, mikrokontroler ESP32, aktuator berupa heater dan water pump, serta antarmuka web untuk visualisasi real-time. Data dikirim dari node sensor ke server melalui protokol MQTT, lalu ditampilkan ke pengguna melalui website berbasis HTTP. Pengujian dilakukan dalam kondisi simulasi dan real-time untuk membandingkan respons kedua metode terhadap perubahan suhu dan turbidity.")
p("")
p("Hasil pengujian menunjukkan bahwa Fuzzy Logic memberikan respons yang lebih stabil dengan overshoot lebih rendah dan steady-state error yang lebih kecil, meskipun memiliki settling time sedikit lebih lama dibanding PID Controller. Sebaliknya, PID Controller memberikan respon lebih cepat namun cenderung mengalami osilasi di sekitar setpoint. Sistem yang dibangun berhasil memantau dan mengontrol kualitas air secara otomatis dan real-time, serta memberikan rekomendasi metode kontrol yang optimal berdasarkan kebutuhan pengguna.")
p("")
p("Kata Kunci: PID Controller, Fuzzy Logic, IoT, monitoring kualitas air, akuarium, ESP32, sensor turbidity AB147")
doc.add_page_break()

# ========================
# BAB I – PENDAHULUAN
# ========================
h1("BAB I – PENDAHULUAN")

h2("1.1 Latar Belakang")
p("Indonesia, sebagai negara kepulauan terbesar di dunia, memiliki kekayaan biodiversitas laut yang sangat tinggi. Keindahan terumbu karang, ikan hias, dan biota laut lainnya telah menjadikan akuarium sebagai salah satu sarana untuk memperkenalkan dan melestarikan keanekaragaman hayati laut kepada masyarakat...")
p("Akuarium ikan hias membutuhkan perhatian khusus dalam menjaga kualitas air agar sesuai dengan kebutuhan spesifik berbagai spesies ikan dan tanaman akuatik. Kualitas air yang buruk dapat menyebabkan stres, penyakit, dan bahkan kematian pada ikan, yang sering kali disebabkan oleh fluktuasi suhu, kekeruhan (turbidity), tingkat oksigen terlarut, dan parameter kimia lainnya (Gawad & Hammad, 2021).")
p("Dalam sistem kendali otomatis, PID Controller merupakan salah satu teknik pengendalian yang paling sering digunakan... Di sisi lain, Fuzzy Logic dianggap mampu memetakan input ke output tanpa mengabaikan ketidakpastian dalam data...")

h2("1.2 Teori")

# IoT
h2("Internet of Things (IoT)")
p("Internet of Things (IoT) adalah paradigma komputasi modern yang menghubungkan objek fisik—seperti sensor, aktuator, dan perangkat elektronik—ke internet untuk memungkinkan pengumpulan, pertukaran, dan analisis data secara otomatis...")
img("gambar1_iot.png")
caption("Gambar 1. Arsitektur Sistem IoT pada Monitoring Kualitas Air Akuarium")
doc.add_paragraph()

# Sensor Suhu
h2("Sensor Suhu DS18B20")
p("Sensor DS18B20 adalah sensor suhu digital berbasis protokol 1-Wire yang mampu mengukur suhu dalam rentang –55°C hingga +125°C dengan akurasi ±0.5°C...")

# Sensor Turbidity AB147 (PENGANTIAN UTAMA)
h2("Sensor Turbidity AB147")
p("Sensor turbidity AB147 (Gravity: Analog Turbidity Sensor) adalah modul sensor analog yang dirancang untuk mengukur tingkat kekeruhan (turbidity) dalam air berdasarkan prinsip nephelometric light scattering. Sensor ini bekerja dengan memancarkan cahaya inframerah melalui sampel air dan mengukur intensitas cahaya yang dipantulkan oleh partikel tersuspensi (seperti sisa pakan, kotoran ikan, atau mikroorganisme).")
p("Sensor ini menghasilkan output tegangan analog 0–5 V, di mana tegangan rendah menunjukkan air jernih (turbidity rendah) dan tegangan tinggi menunjukkan air keruh (turbidity tinggi). Rentang pengukuran sensor AB147 adalah 0–1000 NTU dengan akurasi ±5%.")
p("Dalam konteks akuarium, turbidity ideal harus mendekati 0–5 NTU. Air yang keruh dapat mengurangi penetrasi cahaya, menjadi indikator akumulasi limbah organik, dan menyebabkan stres pada ikan. Oleh karena itu, sistem otomatis akan mengaktifkan water pump atau sistem filtrasi ketika turbidity melebihi ambang batas.")
doc.add_paragraph()

# ESP32, PID, Fuzzy, MQTT, Parameter Evaluasi — tetap sama
h2("Mikrokontroler ESP32")
p("ESP32 adalah mikrokontroler berbasis Wi-Fi dan Bluetooth dual-core...")

h2("PID Controller")
p("PID (Proportional-Integral-Derivative) Controller adalah metode kendali umpan balik...")
p("Rumus dasar: u(t) = Kₚ·e(t) + Kᵢ∫e(τ)dτ + K_d·de(t)/dt")
img("gambar2_pid.png")
caption("Gambar 2. Kurva Respons PID Controller terhadap Gangguan Suhu")
doc.add_paragraph()

h2("Fuzzy Logic Controller")
p("Fuzzy Logic adalah pendekatan kendali berbasis logika linguistik...")
img("gambar3_fuzzy.png")
caption("Gambar 3. Blok Diagram Fuzzy Logic Controller")
doc.add_paragraph()

h2("Aktuator: Heater dan Water Pump")
p("1. Heater (Pemanas Air): Heater 50W digunakan untuk meningkatkan suhu air...\n2. Water Pump (Pompa Air): Water pump 12V berfungsi untuk mengalirkan air bersih dari reservoir atau mengaktifkan sistem filtrasi ketika sensor turbidity AB147 mendeteksi kekeruhan di atas ambang batas (misalnya >10 NTU).")
doc.add_paragraph()

h2("Protokol MQTT")
p("MQTT adalah protokol komunikasi ringan berbasis publish-subscribe...")

h2("Parameter Evaluasi Sistem Kendali")
p("Untuk membandingkan performa PID Controller dan Fuzzy Logic secara objektif, tiga parameter teknis utama digunakan:\n1. Overshoot\n2. Steady-State Error\n3. Settling Time")
img("gambar4_eval.png")
caption("Gambar 4. Parameter Evaluasi Sistem Kendali")

# Rumusan Masalah, Tujuan, Manfaat, Ruang Lingkup
h2("1.3 Rumusan Masalah")
p("1. Bagaimana merancang sistem pengontrol kualitas air dalam akuarium dengan menerapkan PID Controller dan Fuzzy Logic?\n2. Bagaimana perbandingan performa antara PID Controller dan Fuzzy Logic dalam memonitoring kualitas air akuarium berdasarkan parameter overshoot, steady-state error, dan settling time?")

h2("1.4 Tujuan Penelitian")
p("1. Merancang dan membangun sistem monitoring kualitas air pada akuarium dengan menerapkan PID Controller dan Fuzzy Logic.\n2. Menunjukkan perbandingan performa antara PID Controller dan Fuzzy Logic dalam memonitoring kualitas air akuarium berdasarkan parameter teknis: overshoot, steady-state error, dan settling time.")

h2("1.5 Manfaat Penelitian")
p("1. Menjadi acuan dalam merancang sistem monitoring kualitas air akuarium secara real-time berbasis IoT.\n2. Membantu pengguna dalam menentukan metode kontrol yang paling optimal.\n3. Memberikan kontribusi ilmiah dalam penerapan metode kendali cerdas pada sistem IoT bidang akuakultur.")

h2("1.6 Ruang Lingkup")
p("1. Sistem hanya mengontrol dua parameter utama: suhu dan kekeruhan (turbidity) air.\n2. Sensor yang digunakan: DS18B20 (suhu) dan AB147 (turbidity).\n3. Mikrokontroler: ESP32.\n4. Aktuator: heater (untuk suhu) dan water pump (untuk turbidity).\n5. Visualisasi data dilakukan melalui website.\n6. Komunikasi antar node menggunakan protokol MQTT.\n7. Parameter evaluasi performa: overshoot, steady-state error, dan settling time.\n8. Pengujian dilakukan di lingkungan laboratorium.")

# ========================
# BAB II – METODOLOGI
# ========================
h1("BAB II – METODOLOGI PENELITIAN")

h2("2.1 Waktu dan Lokasi Penelitian")
p("Penelitian ini dilaksanakan mulai November 2024 hingga Oktober 2025 di Laboratorium Ubiquitous Cloud Computing...")

h2("2.2 Instrumen Penelitian")
p("Tabel 2.1 Spesifikasi Perangkat Keras")
p("- Mikrokontroler: ESP32 DevKit V1\n- Sensor Suhu: DS18B20\n- Sensor Turbidity: AB147\n- Aktuator: Heater 50W, Water Pump 12V")

p("Tabel 2.2 Spesifikasi Perangkat Lunak")
p("- OS: Windows 11\n- Bahasa: C++, Python, JavaScript\n- Library: OneWire, DallasTemperature, PubSubClient, scikit-fuzzy")

h2("2.3 Tahapan Penelitian")
p("1. Studi Literatur\n2. Analisis Kebutuhan\n3. Perancangan Sistem\n4. Implementasi Sistem\n5. Pengujian dan Pengambilan Data\n6. Analisis Performa\n7. Kesimpulan dan Dokumentasi")

h2("2.4 Teknik Pengumpulan Data")
p("Data dikumpulkan melalui sensor node (suhu & turbidity setiap 5 detik), aktuator node, server (MySQL), dan website (grafik real-time).")

h2("2.5 Perancangan Sistem")
p("Sistem terdiri dari tiga lapisan: Edge (ESP32 + sensor/aktuator), Network (Wi-Fi → MQTT), Application (Flask + website).")

h2("2.6 Implementasi Sistem")
p("PID: Kp=25, Ki=1.5, Kd=4\nFuzzy: input error & delta error, membership triangular, inferensi Mamdani, defuzzifikasi centroid.")

h2("2.7 Evaluasi Sistem")
p("Evaluasi berdasarkan overshoot, steady-state error, settling time. Pengujian: gangguan suhu (es) dan turbidity (partikel), 5 kali ulangan.")

# ========================
# BAB III – HASIL & PEMBAHASAN
# ========================
h1("BAB III – HASIL DAN PEMBAHASAN")

h2("3.1 Hasil Penelitian")
p("Sistem berhasil memantau suhu dan turbidity secara real-time. Arsitektur: ESP32 → MQTT → Flask → Website.")

h2("3.1.4 Hasil Pengujian Respons Sistem")
p("Pengujian gangguan suhu (es → 24°C) dan turbidity (partikel → 50 NTU):")
p("- PID: overshoot ~3.5% (suhu), ~5% (turbidity), settling time ~75 detik")
p("- Fuzzy: overshoot ~0.8% (suhu), ~1.2% (turbidity), settling time ~105 detik")

h2("3.1.5 Data Kuantitatif Performa")
p("Tabel Performa:")
p("Parameter          | PID Controller | Fuzzy Logic")
p("Overshoot (%)      | 3.5% / 5%      | 0.8% / 1.2%")
p("Steady-state error | ±0.25°C / ±2 NTU | ±0.10°C / ±0.5 NTU")
p("Settling time (s)  | 75             | 105")

h2("3.2 Pembahasan")
p("Fuzzy Logic lebih stabil dan akurat, cocok untuk akuarium rumahan. PID lebih cepat, cocok untuk sistem terkontrol.")

# ========================
# BAB IV – KESIMPULAN & SARAN
# ========================
h1("BAB IV – KESIMPULAN DAN SARAN")

h2("4.1 Kesimpulan")
p("1. Sistem berhasil dibangun dan berfungsi real-time.\n2. Fuzzy Logic: stabil, akurat, aman.\n3. PID Controller: cepat, agresif.\n4. Dashboard web mendukung perbandingan dinamis.")

h2("4.2 Saran")
p("1. Integrasi kalibrasi otomatis sensor AB147.\n2. Penyesuaian parameter PID via MQTT.\n3. Tambah notifikasi (Telegram/email).\n4. Pengujian jangka panjang di akuarium nyata.")

# ========================
# SIMPAN FILE
# ========================
doc.save("SkripsiBelumFix.docx")
print("✅ File 'Skripsi_Hafiz_Turbidity_AB147.docx' berhasil dibuat!")
print("📁 4 file placeholder gambar juga telah dibuat.")